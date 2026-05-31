'''
Multi-format support (PDF, DOCX, TXT, MD, HTML)
- Intelligent chunking with overlap
- Metadata extraction (title, author, date, file type)
- Text cleaning and normalization
- Duplicate detection
'''
import hashlib
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Protocol

import PyPDF2
import tiktoken
from bs4 import BeautifulSoup
from docx import Document


class _RegexTokenizer:
    """Offline fallback tokenizer when tiktoken encoding cannot be loaded."""

    _token_pattern = re.compile(r"\w+|[^\w\s]", re.UNICODE)

    def encode(self, text: str) -> List[str]:
        return self._token_pattern.findall(text)

    def decode(self, token_ids: List[str]) -> str:
        if not token_ids:
            return ""
        return " ".join(token_ids)


class _Tokenizer(Protocol):
    def encode(self, text: str) -> List[Any]:
        ...

    def decode(self, token_ids: List[Any]) -> str:
        ...


class DocumentProcessor:
    def __init__(
        self,
        chunk_size: int = 600,
        overlap: int = 100,
        tokenizer_name: str = "gpt2",
        chunk_strategy: str = "tiktoken",
    ):
        if chunk_size <= 0:
            raise ValueError("chunk_size must be > 0")
        if overlap < 0:
            raise ValueError("overlap must be >= 0")
        if overlap >= chunk_size:
            raise ValueError("overlap must be smaller than chunk_size")
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.tokenizer_name = tokenizer_name
        self.chunk_strategy = (chunk_strategy or "tiktoken").strip().lower()
        self._tokenizer: _Tokenizer
        try:
            self._tokenizer = tiktoken.get_encoding(tokenizer_name)
        except Exception:
            self._tokenizer = _RegexTokenizer()
        self._seen_hashes: set = set()

    def process_document(self, file_path: str) -> Optional[Dict]:
        text = self.extract_text(file_path)
        if self._is_duplicate(text):
            return None
        metadata = self.extract_metadata(file_path)
        cleaned_text = self.clean_text(text)
        chunks = self.chunk_text(cleaned_text)
        return {
            'metadata': metadata,
            'chunks': chunks,
        }

    def _is_duplicate(self, text: str) -> bool:
        digest = hashlib.sha256(text.encode('utf-8')).hexdigest()
        if digest in self._seen_hashes:
            return True
        self._seen_hashes.add(digest)
        return False

    def extract_text(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        extractors = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.txt': self._extract_plain_text,
            '.md': self._extract_plain_text,
            '.html': self._extract_html_text,
        }
        extractor = extractors.get(ext)
        if extractor is None:
            raise ValueError(f"Unsupported file type: {ext!r}")
        return extractor(file_path)

    def extract_metadata(self, file_path: str) -> Dict:
        ext = os.path.splitext(file_path)[1].lower()
        base = {
            'title': os.path.basename(file_path),
            'author': 'Unknown',
            'date': None,
            'file_type': ext,
        }
        if ext == '.pdf':
            base.update(self._pdf_metadata(file_path))
        elif ext == '.docx':
            base.update(self._docx_metadata(file_path))
        if base['date'] is None:
            base['date'] = datetime.now().isoformat()
        return base

    def clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def chunk_text(self, text: str) -> List[str]:
        strategy = self.chunk_strategy
        if strategy == "tiktoken":
            return self._chunk_text_token_window(text)
        if strategy == "spacy":
            return self._chunk_text_spacy(text)
        if strategy == "nltk":
            return self._chunk_text_nltk(text)
        if strategy == "medical":
            return self._chunk_text_domain(text, domain="medical")
        if strategy == "legal":
            return self._chunk_text_domain(text, domain="legal")
        raise ValueError(f"Unknown chunking strategy: {strategy!r}")

    def _chunk_text_token_window(self, text: str) -> List[str]:
        token_ids = self._tokenizer.encode(text)
        if not token_ids:
            return []

        chunks: List[str] = []
        step = self.chunk_size - self.overlap
        start = 0
        while start < len(token_ids):
            end = min(start + self.chunk_size, len(token_ids))
            chunk = self._tokenizer.decode(token_ids[start:end])
            if chunk.strip():
                chunks.append(chunk)
            start += step
        return chunks

    def _chunk_text_spacy(self, text: str) -> List[str]:
        sentences = self._sentences_spacy(text)
        return self._chunk_sentences(sentences)

    def _sentences_spacy(self, text: str) -> List[str]:
        try:
            import spacy  # noqa: PLC0415
        except Exception as exc:  # pragma: no cover - depends on optional runtime deps
            raise ValueError("spaCy is not installed. Install spaCy and an English model.") from exc

        nlp = None
        load_errors: List[str] = []
        for model_name in ("en_core_web_sm", "en_core_web_md", "en_core_web_lg"):
            try:
                nlp = spacy.load(model_name, disable=["ner", "lemmatizer", "textcat"])
                break
            except Exception as exc:  # pragma: no cover - environment dependent
                load_errors.append(f"{model_name}: {exc}")
        if nlp is None:
            try:
                nlp = spacy.blank("en")
                nlp.add_pipe("sentencizer")
            except Exception as exc:  # pragma: no cover
                details = "; ".join(load_errors)
                raise ValueError(f"spaCy sentence pipeline unavailable. {details}") from exc

        doc = nlp(text)
        return [s.text.strip() for s in doc.sents if s.text.strip()]

    def _chunk_text_nltk(self, text: str) -> List[str]:
        sentences = self._sentences_nltk(text)
        return self._chunk_sentences(sentences)

    def _sentences_nltk(self, text: str) -> List[str]:
        try:
            import nltk  # noqa: PLC0415
            from nltk.tokenize import sent_tokenize  # noqa: PLC0415
        except Exception as exc:  # pragma: no cover
            raise ValueError("NLTK is not installed. Install nltk package.") from exc

        try:
            sentences = sent_tokenize(text)
        except LookupError:
            nltk.download("punkt", quiet=True)
            try:
                nltk.download("punkt_tab", quiet=True)
            except Exception:
                pass
            sentences = sent_tokenize(text)
        return [s.strip() for s in sentences if s.strip()]

    def _chunk_text_domain(self, text: str, *, domain: str) -> List[str]:
        if domain == "medical":
            boundaries = re.compile(
                (
                    r"(?i)\b(history of present illness|assessment and plan|chief complaint|"
                    r"diagnosis|medications|allergies|impression|plan)\b"
                )
            )
        else:
            boundaries = re.compile(
                r"(?i)\b(section\s+\d+|article\s+\d+|clause\s+\d+|whereas|hereby|pursuant to|party|agreement)\b"
            )

        sentences = self._split_sentences_with_nlp_fallback(text)
        if not sentences:
            return self._chunk_text_token_window(text)
        units: List[str] = []
        current: List[str] = []
        for sentence in sentences:
            if boundaries.search(sentence) and current:
                units.append(" ".join(current).strip())
                current = [sentence]
            else:
                current.append(sentence)
        if current:
            units.append(" ".join(current).strip())
        return self._chunk_sentences(units)

    def _split_sentences_with_nlp_fallback(self, text: str) -> List[str]:
        try:
            sentences = self._sentences_spacy(text)
            if sentences:
                return sentences
        except Exception:
            pass
        try:
            sentences = self._sentences_nltk(text)
            if sentences:
                return sentences
        except Exception:
            pass
        try:
            return [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
        except Exception:
            return [text.strip()] if text.strip() else []

    def _chunk_sentences(self, sentences: List[str]) -> List[str]:
        if not sentences:
            return []
        chunks: List[str] = []
        current: List[str] = []
        current_tokens = 0
        overlap_tail: List[str] = []

        for sentence in sentences:
            sent_tokens = max(1, self.count_tokens(sentence))
            if current and (current_tokens + sent_tokens) > self.chunk_size:
                chunk_text = " ".join(current).strip()
                if chunk_text:
                    chunks.append(chunk_text)
                overlap_tail = self._overlap_tail_sentences(current)
                current = list(overlap_tail)
                current_tokens = sum(max(1, self.count_tokens(s)) for s in current)
            current.append(sentence)
            current_tokens += sent_tokens

        final_chunk = " ".join(current).strip()
        if final_chunk:
            chunks.append(final_chunk)
        return chunks

    def _overlap_tail_sentences(self, sentences: List[str]) -> List[str]:
        if self.overlap <= 0 or not sentences:
            return []
        tail: List[str] = []
        token_count = 0
        for sentence in reversed(sentences):
            s_tokens = max(1, self.count_tokens(sentence))
            if token_count + s_tokens > self.overlap and tail:
                break
            tail.append(sentence)
            token_count += s_tokens
            if token_count >= self.overlap:
                break
        tail.reverse()
        return tail

    def count_tokens(self, text: str) -> int:
        return len(self._tokenizer.encode(text))

    # --- private extractors ---

    def _extract_pdf_text(self, file_path: str) -> str:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return ''.join(page.extract_text() or '' for page in reader.pages)

    def _extract_docx_text(self, file_path: str) -> str:
        doc = Document(file_path)
        return '\n'.join(para.text for para in doc.paragraphs)

    def _extract_plain_text(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _extract_html_text(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
        return soup.get_text(separator=' ')

    # --- private metadata helpers ---

    def _pdf_metadata(self, file_path: str) -> Dict:
        result = {}
        try:
            with open(file_path, 'rb') as f:
                info: Dict = dict(PyPDF2.PdfReader(f).metadata or {})
            if info.get('/Title'):
                result['title'] = info['/Title']
            if info.get('/Author'):
                result['author'] = info['/Author']
            if info.get('/CreationDate'):
                result['date'] = info['/CreationDate']
        except Exception:
            pass
        return result

    def _docx_metadata(self, file_path: str) -> Dict:
        result = {}
        try:
            props = Document(file_path).core_properties
            if props.title:
                result['title'] = props.title
            if props.author:
                result['author'] = props.author
            if props.created:
                result['date'] = props.created.isoformat()
        except Exception:
            pass
        return result
